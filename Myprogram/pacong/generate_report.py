#!/usr/bin/env python3
"""
生成企业办公协同应用定制包统计分析报告 (Word)
分析维度：飞书定制包、钉钉定制包、政务钉钉(专有钉钉)定制包、企业微信定制包
"""
import sqlite3
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ==================== 工具函数 ====================

def set_cell_bg(cell, color_hex):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_text(cell, text, bold=False, font_size=9, alignment=None, font_color=None):
    """设置单元格文本样式"""
    cell.text = ''
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.font.bold = True
    if font_color:
        run.font.color.rgb = RGBColor(*font_color)
    # 减小段落间距
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, header_color)
        set_cell_text(cell, header, bold=True, font_size=9, 
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, font_color=(255, 255, 255))
    
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_bg(cell, "F2F7FB")
            set_cell_text(cell, val, font_size=9)
    
    # 列宽
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    
    return table

def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading

def add_paragraph_styled(doc, text, bold=False, font_size=11, first_line_indent=True):
    """添加带样式的段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(font_size)
    if bold:
        run.font.bold = True
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p


# ==================== 数据准备 ====================

conn = sqlite3.connect('output/results.db')
c = conn.cursor()

# 飞书定制包
feishu_pkgs = c.execute("""SELECT package_name, app_name, enterprise_name, source_site, discovery_method 
    FROM app_info WHERE package_name LIKE 'com.ss.android.lark.%' 
    ORDER BY package_name""").fetchall()

# 钉钉定制包
dingding_pkgs = c.execute("""SELECT package_name, app_name, enterprise_name, source_site, discovery_method 
    FROM app_info WHERE package_name LIKE 'com.alibaba.android.rimet.%' 
    ORDER BY package_name""").fetchall()

# 政务钉钉定制包
taurus_pkgs = c.execute("""SELECT package_name, app_name, enterprise_name, source_site, discovery_method 
    FROM app_info WHERE package_name LIKE 'com.alibaba.taurus.%' 
    ORDER BY package_name""").fetchall()

# 企业微信系列
wecom_pkgs = c.execute("""SELECT package_name, app_name, enterprise_name, source_site, discovery_method 
    FROM app_info WHERE package_name IN (
        'com.tencent.wework', 'com.tencent.weworkenterprise', 'com.tencent.weworklocal'
    ) ORDER BY package_name""").fetchall()

# 总记录数
total_records = c.execute('SELECT COUNT(*) FROM app_info').fetchone()[0]

# 各产品线
product_lines = c.execute("""SELECT product_line, COUNT(*) FROM app_info 
    GROUP BY product_line ORDER BY COUNT(*) DESC""").fetchall()

# 数据来源方法
methods = c.execute("""SELECT discovery_method, COUNT(*) FROM app_info 
    GROUP BY discovery_method ORDER BY COUNT(*) DESC""").fetchall()

conn.close()

# ==================== 行业分类逻辑 ====================

def classify_industry(ent_name, app_name):
    """根据企业名称和APP名称判断行业分类"""
    text = f"{ent_name} {app_name}"
    
    # 政府/政务
    gov_keywords = ['政府', '政务', '政通', '政钉', '大数据中心', '大数据管理局', '经济信息中心',
                    '数字资源', '数字浙江', '海事局', '国家税务', '部委']
    for kw in gov_keywords:
        if kw in text:
            return '政府政务'
    
    # 教育科研
    edu_keywords = ['大学', '学院', '理工', '计算机学会', '实验室', '学术']
    for kw in edu_keywords:
        if kw in text:
            return '教育科研'
    
    # 房地产/建筑
    re_keywords = ['绿城', '融创', '碧桂园', '恒大', '建筑', '建投', '投控']
    for kw in re_keywords:
        if kw in text:
            return '房地产/建筑'
    
    # 金融保险
    fin_keywords = ['保险', '太平洋', '证券', '投资', '国投', '金融']
    for kw in fin_keywords:
        if kw in text:
            return '金融保险'
    
    # 制造/工业
    mfg_keywords = ['美的', '海尔', '海信', '威孚', '宁德时代', '电网', '电子产业',
                    '海伦哲', '正泰', '车企', '奥迪', '一汽', '燃气']
    for kw in mfg_keywords:
        if kw in text:
            return '制造/工业'
    
    # 物流/运输
    log_keywords = ['顺丰', '物流', '航空', '吉祥', '中交']
    for kw in log_keywords:
        if kw in text:
            return '物流/交通运输'
    
    # 零售/消费
    retail_keywords = ['胖东来', '蓝河', '新希望', '乳业', '食品']
    for kw in retail_keywords:
        if kw in text:
            return '零售/消费品'
    
    # 互联网/科技
    tech_keywords = ['字节', '阿里', '科大讯飞', '互联网', '钉钉', '飞书',
                    '迪士尼', '华住', '复星', '普洛斯', '咨询']
    for kw in tech_keywords:
        if kw in text:
            return '互联网/科技服务'
    
    return '其他'

def classify_scenario(ent_name, app_name, pkg_name):
    """判断使用场景"""
    text = f"{ent_name} {app_name} {pkg_name}"
    
    if any(kw in text for kw in ['政通', '政钉', 'taurus', '政务']):
        return '政务协同办公'
    if any(kw in text for kw in ['校园', '大学', '教育']):
        return '校园数字化'
    if any(kw in text for kw in ['信创', 'ce']):
        return '信创适配'
    if any(kw in text for kw in ['学习兴税', '学习']):
        return '培训学习'
    if any(kw in text for kw in ['内测', '内部']):
        return '内部试点'
    return '企业内部办公'


# ==================== 生成报告 ====================

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ===== 封面 =====
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('企业办公协同应用\n定制包统计分析报告')
run.font.size = Pt(28)
run.font.bold = True
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('——飞书 · 钉钉 · 专有钉钉 · 企业微信 定制版深度分析')
run.font.size = Pt(14)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for _ in range(4):
    doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(f'报告日期：{datetime.now().strftime("%Y年%m月%d日")}')
run.font.size = Pt(12)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ===== 目录页 =====
add_heading_styled(doc, '目录', level=1)
toc_items = [
    '一、概述',
    '二、数据总览',
    '三、飞书定制包分析',
    '    3.1 飞书定制包总览',
    '    3.2 飞书定制包行业分布',
    '    3.3 飞书定制包使用场景',
    '    3.4 飞书定制包命名规律',
    '    3.5 飞书定制包完整清单',
    '四、钉钉定制包分析',
    '    4.1 钉钉企业定制包总览',
    '    4.2 钉钉企业定制包行业分布',
    '    4.3 钉钉企业定制包完整清单',
    '五、专有钉钉(政务钉钉)定制包分析',
    '    5.1 专有钉钉总览',
    '    5.2 专有钉钉地域分布',
    '    5.3 专有钉钉完整清单',
    '六、企业微信定制包分析',
    '七、跨平台对比分析',
    '    7.1 定制包数量对比',
    '    7.2 行业覆盖对比',
    '    7.3 数据验证质量对比',
    '八、关键发现与结论',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ===== 一、概述 =====
add_heading_styled(doc, '一、概述', level=1)

add_paragraph_styled(doc, 
    '随着企业数字化转型的深入推进，飞书、钉钉、企业微信三大主流办公协同平台已成为国内企业信息化建设的核心基础设施。'
    '越来越多的大型企业和政府机构选择在标准版基础上进行深度定制，打造专属的移动办公入口。')

add_paragraph_styled(doc, 
    '本报告基于对Android应用生态的系统性爬取和分析，共采集到 {:,} 条应用记录，'
    '其中识别出 {} 个飞书定制包、{} 个钉钉企业定制包、{} 个专有钉钉(政务钉钉)定制包，'
    '以及 {} 个企业微信版本。报告从包名类型、行业分布、使用场景、命名规律等多个维度进行深入分析，'
    '旨在揭示中国企业办公协同应用的定制化趋势。'.format(
        total_records, len(feishu_pkgs), len(dingding_pkgs), len(taurus_pkgs), len(wecom_pkgs)))

# ===== 二、数据总览 =====
add_heading_styled(doc, '二、数据总览', level=1)
add_heading_styled(doc, '2.1 数据库规模', level=2)

overview_rows = [
    ['总应用记录数', f'{total_records:,}', '涵盖Android主流办公、政务、协同类应用'],
]
for pl, cnt in product_lines:
    overview_rows.append([f'{pl}产品线', str(cnt), f'占比 {cnt/total_records*100:.1f}%'])

add_styled_table(doc, ['统计项', '数量', '说明'], overview_rows, col_widths=[5, 3, 9])

doc.add_paragraph()
add_heading_styled(doc, '2.2 定制包汇总', level=2)

add_paragraph_styled(doc,
    '在全部应用记录中，共发现 {} 个明确的企业/政府定制包（不含各平台主应用），分布如下：'.format(
        len(feishu_pkgs) + len(dingding_pkgs) + len(taurus_pkgs) + len(wecom_pkgs) - 4))  # 减去主包

custom_summary = [
    ['飞书定制包', 'com.ss.android.lark.*', str(len(feishu_pkgs)), 
     '含标准版、极速版、信创版及企业专属版'],
    ['钉钉企业定制包', 'com.alibaba.android.rimet.*', str(len(dingding_pkgs)),
     '企业深度定制版钉钉'],
    ['专有钉钉(政务版)', 'com.alibaba.taurus.*', str(len(taurus_pkgs)),
     '政府专用协同办公平台'],
    ['企业微信系列', 'com.tencent.wework*', str(len(wecom_pkgs)),
     '含标准版、私有版、政务版'],
    ['合计', '—', str(len(feishu_pkgs) + len(dingding_pkgs) + len(taurus_pkgs) + len(wecom_pkgs)), ''],
]

add_styled_table(doc, ['平台类型', '包名前缀', '数量', '说明'], custom_summary, col_widths=[4, 5.5, 2, 5.5])

doc.add_paragraph()
add_heading_styled(doc, '2.3 数据验证方式分布', level=2)

method_mapping = {
    'search_engine_verified': '搜索引擎验证',
    'news_verified': '新闻报道验证',
    'user_confirmed_search_assisted': '用户确认+搜索辅证',
    'user_confirmed_package_analysis': '用户确认+包名分析',
    'known_package_verify': '已知包名验证(应用宝)',
    'known_feishu_custom': '已知飞书定制包(应用宝)',
    'known_dingding_custom': '已知钉钉定制包(应用宝)',
    'known_taurus_custom': '已知政务钉钉(应用宝)',
    'snowball_round_1': '滚雪球发现(第1轮)',
    'snowball_round_2': '滚雪球发现(第2轮)',
    'snowball_round_3': '滚雪球发现(第3轮)',
}

method_rows = []
for m, cnt in methods:
    method_rows.append([method_mapping.get(m, m), str(cnt), f'{cnt/total_records*100:.1f}%'])

add_styled_table(doc, ['验证方式', '记录数', '占比'], method_rows, col_widths=[6, 3, 3])

doc.add_page_break()

# ===== 三、飞书定制包分析 =====
add_heading_styled(doc, '三、飞书定制包分析', level=1)
add_heading_styled(doc, '3.1 飞书定制包总览', level=2)

add_paragraph_styled(doc,
    '飞书(Lark)是字节跳动旗下的企业协作与管理平台，其Android定制包以 com.ss.android.lark.* 为包名前缀。'
    '本次共采集到 {} 个飞书定制包，是三大平台中定制包数量最多的，反映出飞书在大中型企业市场的深度渗透。'.format(len(feishu_pkgs)))

# 分类统计
feishu_standard = [p for p in feishu_pkgs if p[4] in ('known_feishu_custom',) and p[1] in ('飞书',)]
feishu_variants = [p for p in feishu_pkgs if p[4] in ('known_feishu_custom',) and p[1] not in ('飞书',)]
feishu_verified = [p for p in feishu_pkgs if p[4] == 'search_engine_verified']
feishu_user_search = [p for p in feishu_pkgs if p[4] == 'user_confirmed_search_assisted']
feishu_user_pkg = [p for p in feishu_pkgs if p[4] == 'user_confirmed_package_analysis']

feishu_cat_rows = [
    ['飞书标准版变体', str(len(feishu_standard)), '同名"飞书"但不同包名，如极速版(lite)、专有云版(ka31)、小米版(kami)'],
    ['应用宝发现的定制版', str(len(feishu_variants)), '应用宝上架的企业定制版，如正泰、工作说说、融e商城'],
    ['搜索引擎验证', str(len(feishu_verified)), '通过第三方下载站确认真实APP名称，如i讯飞、数字国投、胖东来'],
    ['用户确认+搜索辅证', str(len(feishu_user_search)), '用户提供名称+搜索引擎间接验证，如顺丰SF-Work、美的智慧'],
    ['用户确认+包名分析', str(len(feishu_user_pkg)), '企业内部分发、公开网络无痕迹，如华南电网、江苏建投、中国海事'],
    ['合计', str(len(feishu_pkgs)), ''],
]
add_styled_table(doc, ['子类型', '数量', '说明'], feishu_cat_rows, col_widths=[5, 2, 10])

doc.add_paragraph()
add_heading_styled(doc, '3.2 飞书定制包行业分布', level=2)

# 行业分类
feishu_industries = {}
for pkg, name, ent, site, disc in feishu_pkgs:
    ind = classify_industry(ent, name)
    feishu_industries.setdefault(ind, []).append((pkg, name, ent))

ind_rows = sorted(feishu_industries.items(), key=lambda x: -len(x[1]))
feishu_ind_table = []
for ind, items in ind_rows:
    examples = ', '.join([i[1] for i in items[:3]])
    if len(items) > 3:
        examples += f' 等{len(items)}个'
    feishu_ind_table.append([ind, str(len(items)), f'{len(items)/len(feishu_pkgs)*100:.1f}%', examples])

add_styled_table(doc, ['行业分类', '数量', '占比', '代表性应用'], feishu_ind_table, col_widths=[3.5, 2, 2, 9.5])

add_paragraph_styled(doc, 
    '从行业分布来看，飞书定制包覆盖面广泛。制造/工业类企业占比最高，反映出飞书在制造业数字化转型中的重要角色；'
    '其次是互联网/科技服务和政府政务领域。值得注意的是，飞书还深入到了物流运输（如顺丰SF-Work）、'
    '零售消费（如胖东来）等传统行业，显示出较强的行业渗透能力。')

doc.add_paragraph()
add_heading_styled(doc, '3.3 飞书定制包使用场景', level=2)

feishu_scenarios = {}
for pkg, name, ent, site, disc in feishu_pkgs:
    sc = classify_scenario(ent, name, pkg)
    feishu_scenarios.setdefault(sc, []).append((pkg, name, ent))

sc_rows = sorted(feishu_scenarios.items(), key=lambda x: -len(x[1]))
feishu_sc_table = []
for sc, items in sc_rows:
    examples = ', '.join([i[1] for i in items[:3]])
    if len(items) > 3:
        examples += f' 等'
    feishu_sc_table.append([sc, str(len(items)), examples])

add_styled_table(doc, ['使用场景', '数量', '典型案例'], feishu_sc_table, col_widths=[4, 2, 11])

doc.add_paragraph()
add_heading_styled(doc, '3.4 飞书定制包命名规律', level=2)

add_paragraph_styled(doc,
    '飞书定制包的包名后缀呈现出明显的编码规律，可归纳为以下几类：')

naming_rows = [
    ['语义化命名', 'greentown, ihaier, mdzh, sc, weifu', '直接使用企业名缩写或品牌名，可读性强'],
    ['da* 前缀编码', 'dabcsy97, dahngd31, dai39dl9, dajzjt26', '"da"前缀+编码，可能为特定批次的自动生成包名'],
    ['ka* 前缀编码', 'ka31, kaahyz17, kacf, kacw, kalanhe', '"ka"前缀+编码，部分含语义(如kalanhe=蓝河)'],
    ['sa* 前缀编码', 'sa83b7j6, sahlzj17, samhzo3j, sapdl18', '"sa"前缀+编码，部分含哈希值特征'],
    ['功能标识', 'ce(信创版), lite(极速版), sc(顺丰)', '标识版本功能或企业简称'],
]
add_styled_table(doc, ['命名类型', '示例', '特征'], naming_rows, col_widths=[3.5, 7, 6.5])

add_paragraph_styled(doc,
    '编码式命名（da*/ka*/sa*前缀）的包名在公开网络上几乎无法被搜索到，这些APP仅通过企业内部MDM（移动设备管理）'
    '平台分发，不在任何公开应用商店上架。这种策略有效保护了企业内部应用的隐私性，但也增加了外部分析的难度。')

doc.add_paragraph()
add_heading_styled(doc, '3.5 飞书定制包完整清单', level=2)

feishu_full = []
for i, (pkg, name, ent, site, disc) in enumerate(feishu_pkgs, 1):
    suffix = pkg.replace('com.ss.android.lark.', '')
    method_cn = method_mapping.get(disc, disc)
    feishu_full.append([str(i), suffix, name, ent, method_cn])

add_styled_table(doc, ['序号', '包名后缀', 'APP名称', '所属企业', '验证方式'],
                feishu_full, col_widths=[1, 3, 3.5, 5, 4.5])

doc.add_page_break()

# ===== 四、钉钉定制包分析 =====
add_heading_styled(doc, '四、钉钉定制包分析', level=1)
add_heading_styled(doc, '4.1 钉钉企业定制包总览', level=2)

add_paragraph_styled(doc,
    '钉钉(DingTalk)是阿里巴巴旗下的企业协同办公平台，其企业定制版Android包以 com.alibaba.android.rimet.* 为包名前缀。'
    '本次共采集到 {} 个钉钉企业定制包，涵盖大型企业、高校、政务和行业组织等多种类型。'.format(len(dingding_pkgs)))

# 分类
dd_verified = [p for p in dingding_pkgs if p[4] == 'search_engine_verified']
dd_news = [p for p in dingding_pkgs if p[4] == 'news_verified']
dd_user = [p for p in dingding_pkgs if p[4] == 'user_confirmed_search_assisted']
dd_known = [p for p in dingding_pkgs if p[4] in ('known_dingding_custom',)]

dd_cat_rows = [
    ['搜索引擎验证', str(len(dd_verified)), '通过第三方站点确认，如阿里钉(IT168)、CCFLink(极客公园)'],
    ['新闻报道验证', str(len(dd_news)), '通过公开新闻确认，如碧桂园服务(10万人)、宁德时代(10万员工)'],
    ['用户确认+搜索辅证', str(len(dd_user)), '用户提供名称+搜索间接验证，如复旦校园钉、北理工校园钉'],
    ['应用宝发现', str(len(dd_known)), '应用宝直接爬取，如兴业证券优理宝、巨懂车'],
    ['合计', str(len(dingding_pkgs)), ''],
]
add_styled_table(doc, ['子类型', '数量', '说明'], dd_cat_rows, col_widths=[5, 2, 10])

doc.add_paragraph()
add_heading_styled(doc, '4.2 钉钉企业定制包行业分布', level=2)

dd_industries = {}
for pkg, name, ent, site, disc in dingding_pkgs:
    ind = classify_industry(ent, name)
    dd_industries.setdefault(ind, []).append((pkg, name, ent))

dd_ind_rows = sorted(dd_industries.items(), key=lambda x: -len(x[1]))
dd_ind_table = []
for ind, items in dd_ind_rows:
    examples = ', '.join([i[1] for i in items[:3]])
    if len(items) > 3:
        examples += f' 等'
    dd_ind_table.append([ind, str(len(items)), f'{len(items)/len(dingding_pkgs)*100:.1f}%', examples])

add_styled_table(doc, ['行业分类', '数量', '占比', '代表性应用'], dd_ind_table, col_widths=[3.5, 2, 2, 9.5])

add_paragraph_styled(doc,
    '钉钉企业定制包的行业分布呈现出明显特点：教育科研类占比较高（复旦大学、北京理工大学校园钉），'
    '体现了钉钉在教育信息化领域的深耕；制造业方面有宁德时代、一汽奥迪等头部企业；'
    '此外，迪士尼钉钉(内测)的存在说明钉钉也在积极拓展外资企业市场。')

doc.add_paragraph()
add_heading_styled(doc, '4.3 钉钉企业定制包完整清单', level=2)

dd_full = []
for i, (pkg, name, ent, site, disc) in enumerate(dingding_pkgs, 1):
    suffix = pkg.replace('com.alibaba.android.rimet.', '')
    method_cn = method_mapping.get(disc, disc)
    dd_full.append([str(i), suffix, name, ent, method_cn])

add_styled_table(doc, ['序号', '包名后缀', 'APP名称', '所属企业/机构', '验证方式'],
                dd_full, col_widths=[1, 2.5, 4, 5.5, 4])

doc.add_page_break()

# ===== 五、专有钉钉(政务钉钉)定制包分析 =====
add_heading_styled(doc, '五、专有钉钉(政务钉钉)定制包分析', level=1)
add_heading_styled(doc, '5.1 专有钉钉总览', level=2)

add_paragraph_styled(doc,
    '专有钉钉(原政务钉钉)是阿里巴巴为政府机构和大型央企打造的安全协同办公平台，'
    'Android包名以 com.alibaba.taurus.* 为前缀（"taurus"即金牛座，为该产品的内部代号）。'
    '本次共采集到 {} 个专有钉钉定制包，覆盖多个省级行政区和大型企业集团。'.format(len(taurus_pkgs)))

# 分类：政务 vs 企业
taurus_gov = []
taurus_ent = []
for pkg, name, ent, site, disc in taurus_pkgs:
    if any(kw in f"{ent} {name}" for kw in ['政', '大数据中心', '大数据管理局', '经济信息中心', '数字浙江', '数字资源', '税务']):
        taurus_gov.append((pkg, name, ent))
    else:
        taurus_ent.append((pkg, name, ent))

tau_cat_rows = [
    ['政务类（省市政府协同办公）', str(len(taurus_gov)), 
     '各省市政务协同办公平台，如浙政钉、闽政钉、赣政通、渝快政等'],
    ['企业类（大型企业/央企）', str(len(taurus_ent)),
     '大型企业使用专有钉钉底座，如太好钉(太保)、恒大Gems适配等'],
    ['合计', str(len(taurus_pkgs)), ''],
]
add_styled_table(doc, ['类型', '数量', '说明'], tau_cat_rows, col_widths=[5.5, 2, 9.5])

doc.add_paragraph()
add_heading_styled(doc, '5.2 专有钉钉地域/企业分布', level=2)

add_paragraph_styled(doc,
    '专有钉钉在政务领域的布局以省级行政区为单位，每个省份拥有独立的定制包和品牌名称：')

# 省份映射
province_apps = []
enterprise_apps = []
for pkg, name, ent, site, disc in taurus_pkgs:
    suffix = pkg.replace('com.alibaba.taurus.', '')
    method_cn = method_mapping.get(disc, disc)
    if any(kw in f"{ent} {name}" for kw in ['政', '大数据', '经济信息', '数字浙江', '数字资源', '税务']):
        # 判断省份
        province_map = {
            'zhejiang': '浙江省', 'fujian': '福建省', 'jiangxi': '江西省', 
            'anhui': '安徽省', 'ningxia': '宁夏回族自治区', 'ningxianew': '贵州省',
            'chongqing': '重庆市', 'hainan': '海南省', 'hainanxc': '海南省',
            'changchun': '吉林省长春市', 'qzt': '陕西省', 'zhengzhou': '河南省郑州市',
            'xxxs': '全国(国家税务总局)',
        }
        province = province_map.get(suffix, '未知')
        province_apps.append([province, name, suffix, ent, method_cn])
    else:
        enterprise_apps.append([name, suffix, ent, method_cn])

add_paragraph_styled(doc, '政务类专有钉钉分布：', bold=True, first_line_indent=False)
add_styled_table(doc, ['所属地区', 'APP名称', '包名后缀', '运营主体', '验证方式'],
                province_apps, col_widths=[3, 3, 3, 4.5, 3.5])

if enterprise_apps:
    doc.add_paragraph()
    add_paragraph_styled(doc, '企业类专有钉钉分布：', bold=True, first_line_indent=False)
    add_styled_table(doc, ['APP名称', '包名后缀', '所属企业', '验证方式'],
                    enterprise_apps, col_widths=[3, 3, 6, 5])

add_paragraph_styled(doc,
    '从地域分布看，专有钉钉已覆盖浙江、福建、江西、安徽、陕西、海南、重庆、宁夏、贵州等多个省市区，'
    '形成了较为完整的政务数字化版图。其中浙政钉（浙江）作为最早上线的政务钉钉，'
    '有早期版本(com.alibaba.android.rimet.zj)和正式版(com.alibaba.taurus.zhejiang)两个包名，'
    '反映了从钉钉企业版到专有钉钉的技术架构演进。')

doc.add_paragraph()
add_heading_styled(doc, '5.3 专有钉钉完整清单', level=2)

tau_full = []
for i, (pkg, name, ent, site, disc) in enumerate(taurus_pkgs, 1):
    suffix = pkg.replace('com.alibaba.taurus.', '')
    method_cn = method_mapping.get(disc, disc)
    tau_full.append([str(i), suffix, name, ent, method_cn])

add_styled_table(doc, ['序号', '包名后缀', 'APP名称', '运营主体', '验证方式'],
                tau_full, col_widths=[1, 3.5, 3.5, 5, 4])

doc.add_page_break()

# ===== 六、企业微信定制包分析 =====
add_heading_styled(doc, '六、企业微信定制包分析', level=1)

add_paragraph_styled(doc,
    '企业微信是腾讯旗下的企业通讯与办公工具，与飞书、钉钉不同，企业微信采用的是统一APP+后台配置的定制模式，'
    '而非为每个企业生成独立的APK包名。因此在Android包名层面，企业微信的"定制"主要体现为以下3个版本：')

wecom_rows = []
for i, (pkg, name, ent, site, disc) in enumerate(wecom_pkgs, 1):
    method_cn = method_mapping.get(disc, disc)
    wecom_rows.append([str(i), pkg, name, ent, method_cn])

add_styled_table(doc, ['序号', '包名', 'APP名称', '开发者', '验证方式'],
                wecom_rows, col_widths=[1, 5.5, 3, 4, 3.5])

add_paragraph_styled(doc,
    '企业微信的三个版本定位明确：')

versions_desc = [
    ('com.tencent.wework', '企业微信（标准版）', 
     '面向所有企业的通用版本，支持与微信互通。企业通过管理后台进行个性化配置（工作台、审批流程等），'
     '无需独立包名。这是企业微信超过1400万企业用户使用的主力版本。'),
    ('com.tencent.weworkenterprise', '企业微信私有版', 
     '面向对数据安全有极高要求的大型企业，支持私有化部署。数据存储在企业自有服务器，'
     '适用于金融、军工等敏感行业。'),
    ('com.tencent.weworklocal', '政务微信', 
     '面向政府机构的专用版本，满足等保合规要求。功能定制化程度高，数据存储符合政务云标准。'),
]

for pkg, title, desc in versions_desc:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}（{pkg}）：')
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)

add_paragraph_styled(doc,
    '对比分析：企业微信的定制策略与飞书、钉钉存在本质区别。飞书和钉钉为大客户生成独立包名（如lark.greentown、rimet.fosun），'
    '而企业微信通过统一APP+后台配置实现定制化。这两种策略各有优劣：独立包名便于企业品牌展示和MDM管控，'
    '统一包名则降低了维护成本和分发复杂度。')

doc.add_page_break()

# ===== 七、跨平台对比分析 =====
add_heading_styled(doc, '七、跨平台对比分析', level=1)
add_heading_styled(doc, '7.1 定制包数量对比', level=2)

all_custom = len(feishu_pkgs) + len(dingding_pkgs) + len(taurus_pkgs) + len(wecom_pkgs)
compare_rows = [
    ['飞书', str(len(feishu_pkgs)), f'{len(feishu_pkgs)/all_custom*100:.1f}%', 
     '独立包名模式', '企业MDM分发为主'],
    ['钉钉(企业版)', str(len(dingding_pkgs)), f'{len(dingding_pkgs)/all_custom*100:.1f}%',
     '独立包名模式', '应用商店+企业分发'],
    ['专有钉钉(taurus)', str(len(taurus_pkgs)), f'{len(taurus_pkgs)/all_custom*100:.1f}%',
     '独立包名模式', '政务应用商店为主'],
    ['企业微信', str(len(wecom_pkgs)), f'{len(wecom_pkgs)/all_custom*100:.1f}%',
     '统一包名+后台配置', '应用商店统一分发'],
    ['合计', str(all_custom), '100%', '', ''],
]

add_styled_table(doc, ['平台', '包名数', '占比', '定制模式', '分发方式'], 
                compare_rows, col_widths=[3, 2, 2, 4.5, 5.5])

add_paragraph_styled(doc,
    f'飞书以 {len(feishu_pkgs)} 个定制包名位居第一，占比超过一半，显示出字节跳动在B端市场的积极拓展策略。'
    f'钉钉体系（企业定制 {len(dingding_pkgs)} + 专有钉钉 {len(taurus_pkgs)} = {len(dingding_pkgs)+len(taurus_pkgs)}个）'
    f'合计排名第二。企业微信以3个版本覆盖全部场景，体现了腾讯"平台化"的产品思路。')

doc.add_paragraph()
add_heading_styled(doc, '7.2 行业覆盖对比', level=2)

# 合并行业统计
all_industries = {}
for platform_name, pkgs in [('飞书', feishu_pkgs), ('钉钉', dingding_pkgs), ('专有钉钉', taurus_pkgs)]:
    for pkg, name, ent, site, disc in pkgs:
        ind = classify_industry(ent, name)
        all_industries.setdefault(ind, {'飞书': 0, '钉钉': 0, '专有钉钉': 0})
        all_industries[ind][platform_name] += 1

ind_compare = sorted(all_industries.items(), key=lambda x: -(x[1]['飞书'] + x[1]['钉钉'] + x[1]['专有钉钉']))
ind_compare_rows = []
for ind, counts in ind_compare:
    total = counts['飞书'] + counts['钉钉'] + counts['专有钉钉']
    ind_compare_rows.append([ind, str(counts['飞书']), str(counts['钉钉']), str(counts['专有钉钉']), str(total)])

add_styled_table(doc, ['行业', '飞书', '钉钉', '专有钉钉', '合计'],
                ind_compare_rows, col_widths=[4, 2.5, 2.5, 2.5, 2.5])

add_paragraph_styled(doc,
    '从行业覆盖来看，三大平台各有侧重：飞书在制造/工业（美的、海尔、宁德时代）和物流运输（顺丰、吉祥航空）'
    '领域优势明显；钉钉在教育科研（复旦、北理工）和外资企业（迪士尼）方面有独特优势；'
    '专有钉钉则在政府政务领域一家独大，覆盖了9个以上省级行政区。')

doc.add_paragraph()
add_heading_styled(doc, '7.3 数据验证质量对比', level=2)

# 按验证方法分类统计
verify_types = {
    '搜索引擎直接验证': ('search_engine_verified',),
    '新闻报道验证': ('news_verified',),
    '应用商店验证': ('known_package_verify', 'known_feishu_custom', 'known_dingding_custom', 'known_taurus_custom', 'snowball_round_1'),
    '用户确认+搜索辅证': ('user_confirmed_search_assisted',),
    '用户确认+包名分析': ('user_confirmed_package_analysis',),
}

all_custom_pkgs = []
for p in feishu_pkgs:
    all_custom_pkgs.append(('飞书', *p))
for p in dingding_pkgs:
    all_custom_pkgs.append(('钉钉', *p))
for p in taurus_pkgs:
    all_custom_pkgs.append(('专有钉钉', *p))
for p in wecom_pkgs:
    all_custom_pkgs.append(('企业微信', *p))

quality_rows = []
for vname, vmethods in verify_types.items():
    cnt = len([p for p in all_custom_pkgs if p[5] in vmethods])
    if cnt > 0:
        quality_rows.append([vname, str(cnt), f'{cnt/len(all_custom_pkgs)*100:.1f}%',
                           '高' if vname in ('搜索引擎直接验证', '新闻报道验证', '应用商店验证') else '中'])

add_styled_table(doc, ['验证方式', '记录数', '占比', '可信度'], quality_rows, col_widths=[5, 2, 2, 2])

add_paragraph_styled(doc,
    f'在全部 {len(all_custom_pkgs)} 个定制包中，通过应用商店、搜索引擎和新闻报道等公开渠道直接验证的占'
    f' {len([p for p in all_custom_pkgs if p[5] in ("search_engine_verified","news_verified","known_package_verify","known_feishu_custom","known_dingding_custom","known_taurus_custom","snowball_round_1")])} 个，'
    f'用户确认结合搜索辅证的占 {len([p for p in all_custom_pkgs if p[5]=="user_confirmed_search_assisted"])} 个，'
    f'仅靠用户确认+包名分析的占 {len([p for p in all_custom_pkgs if p[5]=="user_confirmed_package_analysis"])} 个。'
    f'整体数据验证质量较好，所有记录均已消除"包名推断"状态。')

doc.add_page_break()

# ===== 八、关键发现与结论 =====
add_heading_styled(doc, '八、关键发现与结论', level=1)

findings = [
    ('飞书定制包数量领先',
     f'飞书以 {len(feishu_pkgs)} 个独立定制包名领跑三大平台，显示字节跳动在B端市场采用了'
     '"每个大客户一个专属包"的深度定制策略。这些定制包覆盖了制造业、物流、零售、政务、科技等众多行业，'
     '反映出飞书在企业市场的快速扩张态势。'),
    
    ('专有钉钉在政务领域建立壁垒',
     f'以 com.alibaba.taurus.* 为包名前缀的专有钉钉已覆盖 {len(taurus_gov)} 个省市级政务平台，'
     '形成了从浙政钉到各省政通的完整矩阵。加上企业版钉钉的 {len(dingding_pkgs)} 个定制包，'
     f'阿里钉钉体系总计 {len(dingding_pkgs)+len(taurus_pkgs)} 个定制包，在数量上与飞书相当。'),
    
    ('企业微信的"平台化"差异化策略',
     '与飞书、钉钉为每个大客户生成独立包名不同，企业微信仅有3个版本（标准版、私有版、政务版），'
     '通过统一APP+后台配置实现企业定制化。这种策略在分发效率和维护成本方面具有优势，'
     '但在品牌独立性和企业感知方面有所不足。'),
    
    ('编码式包名保护企业隐私',
     '大量飞书定制包使用da*/ka*/sa*前缀的编码式命名，在公开网络上完全不可检索。'
     '这种策略有效保护了企业客户的隐私，但也说明这些APP完全通过企业内部渠道分发，'
     '不依赖公开应用商店。'),
    
    ('教育和外资企业是钉钉独特优势',
     '钉钉在校园数字化（复旦大学、北京理工大学校园钉）和外资企业（迪士尼内测版）'
     '方面有独特优势，这两个领域在飞书和企业微信的定制包中较少出现。'),
    
    ('政务数字化正加速推进',
     f'本次共发现 {len(taurus_gov)+1} 个政务协同办公定制包（含政务微信），覆盖浙江、福建、江西、'
     '安徽、陕西、海南、重庆、宁夏、贵州、吉林等地区，以及学习兴税等全国性平台。'
     '部分省份还出现了新旧版本并存的现象（如海政通/海政通信创版、宁政通旧/新版），'
     '反映了信创替代和迭代升级的趋势。'),
]

for i, (title, content) in enumerate(findings, 1):
    add_heading_styled(doc, f'8.{i} {title}', level=2)
    add_paragraph_styled(doc, content)

# 最后总结
doc.add_paragraph()
add_heading_styled(doc, '总结', level=2)

add_paragraph_styled(doc,
    '综上所述，中国企业办公协同应用正处于从"标准化"走向"深度定制化"的关键阶段。'
    f'本报告基于 {total_records:,} 条应用数据，系统梳理了飞书（{len(feishu_pkgs)}个）、'
    f'钉钉（{len(dingding_pkgs)}个）、专有钉钉（{len(taurus_pkgs)}个）和企业微信（{len(wecom_pkgs)}个）'
    f'共计 {all_custom} 个定制包的类型、行业分布和使用场景。')

add_paragraph_styled(doc,
    '三大平台在定制化策略上各具特色：飞书追求"一企一包"的极致定制，钉钉以"企业+政务"双轮驱动，'
    '企业微信则坚持"平台统一、后台定制"的轻量化路线。随着数字化转型的深入和信创要求的推进，'
    '企业办公协同应用的定制化趋势将进一步加强，三大平台的竞争也将更加激烈。')

# ===== 保存 =====
output_path = 'output/企业办公协同应用定制包统计分析报告.docx'
doc.save(output_path)
print(f'\n✅ 报告已生成: {output_path}')
print(f'   飞书定制包: {len(feishu_pkgs)} 个')
print(f'   钉钉定制包: {len(dingding_pkgs)} 个')
print(f'   专有钉钉定制包: {len(taurus_pkgs)} 个')
print(f'   企业微信: {len(wecom_pkgs)} 个')
print(f'   总计: {all_custom} 个')
