"""
企业IM安卓客户端包信息爬虫 - 项目开发报告生成脚本
生成完整的项目总结Word文档
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
from datetime import datetime


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading_elm.append(shading)


def set_table_style(table):
    """设置表格样式"""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            if row_idx == 0:
                set_cell_shading(cell, '2B579A')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True


def add_table(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = str(val)
    set_table_style(table)
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)
    doc.add_paragraph('')  # 间距
    return table


def generate_report():
    doc = Document()
    
    # ============================================================
    # 文档样式设置
    # ============================================================
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    for i in range(1, 5):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = '微软雅黑'
        heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        heading_style.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)

    # ============================================================
    # 封面
    # ============================================================
    for _ in range(4):
        doc.add_paragraph('')
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('企业IM安卓客户端包信息爬虫')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('项目开发报告')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x4A, 0x6A, 0x9D)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph('')
    
    # 分隔线
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line_p.add_run('━' * 40)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    run.font.size = Pt(12)

    doc.add_paragraph('')
    
    # 项目信息
    info_items = [
        ('项目名称', '企业IM安卓客户端包信息爬虫'),
        ('项目编号', 'PACONG-2026-001'),
        ('开发日期', '2026年3月9日'),
        ('报告日期', datetime.now().strftime('%Y年%m月%d日')),
        ('文档版本', 'V1.0'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：{value}')
        run.font.size = Pt(13)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_page_break()
    
    # ============================================================
    # 目录页
    # ============================================================
    doc.add_heading('目  录', level=1)
    toc_items = [
        '一、项目概述',
        '二、需求分析',
        '三、系统架构设计',
        '四、技术方案',
        '五、开发过程纪实',
        '六、数据成果',
        '七、项目文件清单',
        '八、问题与解决方案',
        '九、项目总结',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(12)

    doc.add_page_break()
    
    # ============================================================
    # 一、项目概述
    # ============================================================
    doc.add_heading('一、项目概述', level=1)
    
    doc.add_heading('1.1 项目背景', level=2)
    doc.add_paragraph(
        '随着企业数字化转型的深入推进，企业级即时通讯（IM）及协同办公平台在各行业、'
        '各规模企业中的渗透率持续提升。以企业微信、钉钉、飞书为代表的头部产品，不仅'
        '提供标准版本的通用服务，还为大型企业、政府机关等客户提供深度定制的专属版本。'
        '这些定制版以独立安卓应用包（APK）的形式分发，形成了庞大且分散的应用生态。'
    )
    doc.add_paragraph(
        '本项目旨在建立一套自动化的信息采集系统，系统性地发现、验证和记录上述企业IM'
        '产品及其定制版安卓客户端的公开包信息，为企业办公应用市场研究提供数据支撑。'
    )
    
    doc.add_heading('1.2 项目目标', level=2)
    goals = [
        '开发三级递进式搜索爬虫系统，覆盖主流应用商店和搜索引擎',
        '采集企业微信、钉钉、飞书及其定制版客户端的公开信息',
        '建立结构化的应用包信息数据库，支持多维度查询和分析',
        '对采集数据进行质量评分和交叉验证，确保数据可靠性',
        '针对飞书、钉钉、专有钉钉、企业微信等产品线的定制包进行专项统计分析',
    ]
    for g in goals:
        doc.add_paragraph(g, style='List Bullet')
    
    doc.add_heading('1.3 项目范围', level=2)
    add_table(doc, 
        ['维度', '范围说明'],
        [
            ['产品覆盖', '企业微信、钉钉（含专有钉钉/政务钉钉）、飞书及其所有已知定制版'],
            ['数据来源', '应用宝（主要）+ 搜索引擎验证 + 新闻报道确认 + 用户提供信息'],
            ['采集字段', '包名、应用名、企业名、版本号、开发者、下载量、描述等14个字段'],
            ['输出格式', 'SQLite数据库 + CSV/JSON/XLSX多格式导出 + 统计分析Word报告'],
        ],
        col_widths=[3, 13]
    )

    doc.add_page_break()
    
    # ============================================================
    # 二、需求分析
    # ============================================================
    doc.add_heading('二、需求分析', level=1)
    
    doc.add_heading('2.1 功能需求', level=2)
    doc.add_paragraph('项目依据《企业IM安卓客户端包信息爬虫开发需求文档》（pacon_grequire.md）进行开发，核心功能需求如下：')
    
    add_table(doc,
        ['需求编号', '功能模块', '需求描述', '优先级'],
        [
            ['FR-001', '三级递进式搜索', '第一级客户名称挖掘→第二级全网定制包搜索→第三级应用商店补全', 'P0'],
            ['FR-002', '多源数据采集', '支持应用宝、华为、小米、OPPO、vivo等8+应用商店', 'P0'],
            ['FR-003', '搜索引擎爬取', '支持百度、必应、搜狗等搜索引擎的结构化爬取', 'P0'],
            ['FR-004', '数据处理管道', '去重、清洗、企业名提取、产品线分类、质量评分', 'P0'],
            ['FR-005', '多格式导出', 'SQLite/CSV/JSON/XLSX多格式数据导出', 'P1'],
            ['FR-006', '反爬策略', 'UA轮换、请求限速、代理池、重试机制', 'P1'],
            ['FR-007', '定制包验证', '搜索引擎验证、新闻报道确认、用户提供信息交叉验证', 'P0'],
            ['FR-008', '统计分析报告', '按产品线维度生成定制包统计分析Word报告', 'P1'],
        ],
        col_widths=[2.2, 3, 7.5, 1.5]
    )
    
    doc.add_heading('2.2 非功能需求', level=2)
    add_table(doc,
        ['类别', '需求描述'],
        [
            ['合规性', '仅爬取公开可访问页面，遵守robots.txt，不下载APK文件'],
            ['稳定性', '具备完善的异常处理、重试机制，支持断点续爬'],
            ['可扩展性', '模块化架构设计，便于新增应用商店和搜索引擎'],
            ['数据质量', '多源交叉验证，数据质量评分机制，确保采集准确性'],
            ['性能要求', '支持并发请求，单次完整采集耗时控制在合理范围内'],
        ],
        col_widths=[2.5, 13.5]
    )
    
    doc.add_heading('2.3 目标数据字段', level=2)
    add_table(doc,
        ['字段名', '类型', '说明', '示例'],
        [
            ['package_name', 'TEXT', '安卓包名（唯一标识）', 'com.tencent.wework'],
            ['app_name', 'TEXT', '应用显示名称', '企业微信'],
            ['enterprise_name', 'TEXT', '使用/定制该包的企业名称', 'XX集团'],
            ['product_line', 'TEXT', '所属产品线', '企业微信/钉钉/飞书'],
            ['source_site', 'TEXT', '来源网站域名', '应用宝'],
            ['source_url', 'TEXT', '详情页完整URL', 'https://sj.qq.com/...'],
            ['version', 'TEXT', '版本号', '4.1.16'],
            ['download_count', 'TEXT', '下载量', '1000万+'],
            ['description', 'TEXT', '应用简介', '...'],
            ['developer', 'TEXT', '开发者名称', '腾讯科技'],
            ['discovery_method', 'TEXT', '发现方式', 'known_package_verify'],
            ['quality_score', 'REAL', '数据质量评分(0~1)', '0.85'],
            ['category', 'TEXT', '应用分类', '效率办公'],
            ['crawl_time', 'DATETIME', '爬取时间', '2026-03-09 10:30:00'],
        ],
        col_widths=[3, 1.8, 4, 5]
    )

    doc.add_page_break()

    # ============================================================
    # 三、系统架构设计
    # ============================================================
    doc.add_heading('三、系统架构设计', level=1)
    
    doc.add_heading('3.1 整体架构', level=2)
    doc.add_paragraph(
        '系统采用三级递进式搜索架构，配合数据处理管道和统一存储层，形成完整的数据采集、'
        '处理、存储、导出流程。'
    )
    
    # 架构示意文字
    arch_text = (
        '┌─────────────────────────────────────────────────┐\n'
        '│              总调度中心 (main.py)                │\n'
        '│          CrawlerScheduler Pipeline               │\n'
        '└──────┬──────────────┬──────────────┬─────────────┘\n'
        '       │              │              │\n'
        '┌──────▼──────┐ ┌────▼──────┐ ┌─────▼──────┐\n'
        '│ Level1      │ │ Level2    │ │ Level3     │\n'
        '│ 客户名称挖掘│ │ 全网搜索  │ │ 应用商店   │\n'
        '│ (6模块)     │ │ (5模块)   │ │ (9模块)    │\n'
        '└──────┬──────┘ └────┬──────┘ └─────┬──────┘\n'
        '       │              │              │\n'
        '┌──────▼──────────────▼──────────────▼──────┐\n'
        '│        数据处理管道 Pipeline (6模块)        │\n'
        '│  去重 → 清洗 → 企业名提取 → 分类 → 评分   │\n'
        '└──────────────────┬────────────────────────┘\n'
        '                   │\n'
        '┌──────────────────▼────────────────────────┐\n'
        '│          存储与导出层 Storage               │\n'
        '│    SQLite DB + CSV/JSON/XLSX Export         │\n'
        '└───────────────────────────────────────────┘'
    )
    p = doc.add_paragraph()
    run = p.add_run(arch_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    
    doc.add_heading('3.2 模块设计', level=2)
    add_table(doc,
        ['模块', '目录', '文件数', '代码行数', '职责说明'],
        [
            ['第一级爬虫', 'level1_customers/', '6', '1,574', '从官网案例、搜索引擎、新闻报道、政务平台挖掘客户名称'],
            ['第二级爬虫', 'level2_fullweb/', '5', '1,569', '以客户名称为线索，全网搜索定制安卓客户端'],
            ['第三级爬虫', 'level3_appstore/', '9', '1,377', '在8大应用商店中检索补全包信息'],
            ['数据管道', 'pipeline/', '6', '1,204', '去重、清洗、企业名提取、产品线分类、质量评分'],
            ['存储层', 'storage/', '3', '882', '数据库操作封装和多格式导出'],
            ['工具库', 'utils/', '7', '1,181', '搜索引擎封装、HTTP客户端、限速器、UA池'],
            ['业务脚本', '根目录', '42', '8,195', '爬取脚本、诊断工具、定制包处理、报告生成'],
        ],
        col_widths=[2, 3, 1.2, 1.5, 7]
    )
    
    doc.add_heading('3.3 数据流设计', level=2)
    doc.add_paragraph('数据在系统中的流转过程如下：')
    
    flows = [
        '1. 客户名称挖掘 → 输出客户名称清单（customers.db，673条记录）',
        '2. 全网搜索 → 以客户名为线索发现候选定制包',
        '3. 应用商店爬取 → 从应用宝等商店获取完整包信息（含滚雪球发现）',
        '4. 数据处理管道 → 去重、清洗、分类、评分',
        '5. 定制包专项验证 → 搜索引擎验证 + 新闻确认 + 用户提供信息',
        '6. 数据导出 → SQLite/CSV/JSON/XLSX多格式输出',
        '7. 统计分析 → 按产品线维度生成分析报告',
    ]
    for f in flows:
        doc.add_paragraph(f)

    doc.add_page_break()
    
    # ============================================================
    # 四、技术方案
    # ============================================================
    doc.add_heading('四、技术方案', level=1)
    
    doc.add_heading('4.1 技术栈', level=2)
    add_table(doc,
        ['类别', '技术/工具', '版本要求', '用途'],
        [
            ['编程语言', 'Python', '3.11+', '主开发语言'],
            ['爬虫框架', 'Scrapy + Playwright', '≥2.9.0 / ≥1.40.0', '爬虫调度 + JS渲染'],
            ['HTTP客户端', 'aiohttp + httpx', '≥3.9.0 / ≥0.25.0', '异步HTTP请求'],
            ['HTML解析', 'BeautifulSoup4 + lxml', '≥4.12.0 / ≥4.9.0', 'HTML内容解析'],
            ['NLP处理', 'jieba', '≥0.42.1', '中文分词与企业名提取'],
            ['数据处理', 'pandas', '≥2.1.0', '数据处理与导出'],
            ['数据存储', 'SQLite3', '内置', '本地数据库存储'],
            ['文档生成', 'python-docx', '≥0.8.0', 'Word报告生成'],
            ['反爬工具', 'fake-useragent', '≥1.4.0', 'UA轮换'],
            ['去重', 'pybloom-live', '≥4.0.0', 'Bloom Filter高效去重'],
        ],
        col_widths=[2, 3.5, 2.5, 5.5]
    )
    
    doc.add_heading('4.2 爬虫策略', level=2)
    
    doc.add_heading('4.2.1 应用宝爬取（主要数据源）', level=3)
    doc.add_paragraph(
        '应用宝（sj.qq.com）作为主要数据源，采用以下策略：'
    )
    strategies = [
        '种子包名验证：预设120+已知种子包名，逐一在应用宝验证获取详情',
        '滚雪球发现：利用应用宝"猜你喜欢"和"同开发者应用"推荐接口，三轮滚雪球递进扩展',
        '限流应对：降低请求频率（0.5-1.2秒间隔），增加重试机制（MAX_RETRY=2），避免触发限流',
        '数据质量保证：对返回数据进行完整性校验，过滤无效和重复记录',
    ]
    for s in strategies:
        doc.add_paragraph(s, style='List Bullet')
    
    doc.add_heading('4.2.2 搜索引擎验证', level=3)
    doc.add_paragraph(
        '对于应用宝无法直接获取的定制包（企业内部分发、非公开上架），采用搜索引擎验证策略：'
    )
    verifications = [
        '使用web_search工具搜索包名关键词',
        '在第三方APK下载站（火鸟手游网、安粉丝网、当下软件园、IT168、极客公园等）查找详情页',
        '搜索新闻报道间接确认企业定制版的存在',
        '最终成功验证11个包名（搜索引擎直接确认9个 + 新闻报道确认2个）',
    ]
    for v in verifications:
        doc.add_paragraph(v, style='List Bullet')
    
    doc.add_heading('4.2.3 反爬策略', level=3)
    add_table(doc,
        ['策略', '实现方式', '效果'],
        [
            ['请求限速', '每站3~8秒随机延迟，搜索引擎10~15秒', '避免触发限流'],
            ['UA轮换', '50+ User-Agent随机轮换', '模拟正常浏览行为'],
            ['重试机制', '指数退避重试，最多3次', '提升请求成功率'],
            ['Cookie管理', '模拟正常浏览器Cookie', '规避反爬检测'],
            ['并发控制', '最大4个并发请求', '平衡效率与安全'],
        ],
        col_widths=[2.5, 5.5, 5.5]
    )

    doc.add_heading('4.3 数据质量保证', level=2)
    doc.add_paragraph('系统采用多层数据质量保证机制：')
    
    add_table(doc,
        ['机制', '说明', '权重'],
        [
            ['包名完整性', '必须有有效的安卓包名', '25%'],
            ['应用名完整性', '必须有应用显示名称', '15%'],
            ['企业名提取', '能识别出关联企业名称', '15%'],
            ['多源验证', '在多个来源中交叉验证', '15%'],
            ['版本信息', '有版本号信息', '10%'],
            ['开发者信息', '有开发者/发布者信息', '10%'],
            ['描述信息', '有应用描述', '5%'],
            ['下载量', '有下载量数据', '5%'],
        ],
        col_widths=[3, 7, 2]
    )

    doc.add_page_break()
    
    # ============================================================
    # 五、开发过程纪实
    # ============================================================
    doc.add_heading('五、开发过程纪实', level=1)
    
    doc.add_heading('5.1 开发阶段总览', level=2)
    add_table(doc,
        ['阶段', '工作内容', '关键成果'],
        [
            ['阶段一：基础框架搭建', '项目骨架、配置管理、存储层、工具库', '78个Python文件，15,982行代码的模块化架构'],
            ['阶段二：应用宝爬虫(v1/v2)', '初版爬虫开发、问题诊断', '发现应用宝限流问题，仅获得27条记录'],
            ['阶段三：爬虫优化(v3)', '限流应对、滚雪球发现、重试机制', '165条记录（目标≥100 ✅），质量评分0.85'],
            ['阶段四：定制包专项爬取', '71个已知定制包入库', '71/71全部入库，应用宝14个+包名推断54个'],
            ['阶段五：搜索引擎验证', '消除"包名推断"数据', '11个搜索引擎直接确认+2个新闻确认'],
            ['阶段六：用户确认轮次', '根据用户提供APP名称彻底消除推断', '42条推断记录全部更新，包名推断归零'],
            ['阶段七：统计分析报告', '按产品线维度生成定制包分析Word报告', '8章19表的完整分析报告'],
        ],
        col_widths=[3.5, 5, 6]
    )
    
    doc.add_heading('5.2 阶段一：基础框架搭建', level=2)
    doc.add_paragraph(
        '按照需求文档的系统架构设计，搭建了完整的项目骨架，包含7个功能模块共78个Python源文件。'
        '采用模块化设计，各级爬虫、数据管道、存储层、工具库相互独立且通过标准接口协作。'
    )
    doc.add_paragraph('核心模块包括：')
    modules = [
        'level1_customers/：官网案例爬虫、搜索引擎客户挖掘、政务平台爬虫、新闻爬虫等6个模块',
        'level2_fullweb/：搜索引擎包搜索、企业官网扫描、APK站点爬虫等5个模块',
        'level3_appstore/：华为、小米、OPPO、vivo、应用宝、豌豆荚、酷安、APKPure等9个商店爬虫',
        'pipeline/：去重(Bloom Filter)、数据清洗、企业名NLP提取、产品线分类、质量评分等6个处理器',
        'utils/：搜索引擎封装、HTTP客户端、请求限速器、UA池等7个工具',
    ]
    for m in modules:
        doc.add_paragraph(m, style='List Bullet')
    
    doc.add_heading('5.3 阶段二/三：应用宝爬虫优化', level=2)
    doc.add_paragraph('初版爬虫（v2）在实际运行中仅获得27条记录，远低于目标≥100条。通过诊断发现：')
    
    issues = [
        '83个"失败"包名中实际72个可正常获取，仅11个是真正的HTTP 404',
        '失败原因是应用宝限流 —— 请求虽返回HTTP 200，但内容不完整',
    ]
    for i in issues:
        doc.add_paragraph(i, style='List Bullet')
    
    doc.add_paragraph('v3版本的优化措施：')
    optimizations = [
        '清理11个确认404的包名，补充新种子至120个',
        '增加重试机制（MAX_RETRY=2），降低请求频率（0.5-1.2秒间隔）',
        '新增三轮滚雪球发现：利用"猜你喜欢"和"同开发者"推荐接口',
        '最终结果：165条记录，质量评分平均0.85，164条A级',
    ]
    for o in optimizations:
        doc.add_paragraph(o, style='List Bullet')
    
    doc.add_heading('5.4 阶段四：定制包专项爬取', level=2)
    doc.add_paragraph(
        '用户提供71个已知定制安卓包名（飞书40个、钉钉15个、浙政钉16个），要求全部入库。'
    )
    doc.add_paragraph('诊断结果：')
    diagnosis = [
        '71个中仅3个已有记录（飞书主包、钉钉主包、浙政钉浙江版）',
        '应用宝能找到14个（含已有3个），新增11个',
        '其余54个在所有公开应用商店（豌豆荚/酷安/华为/小米）均返回404',
        '原因：企业定制包多通过企业内部MDM分发，不在公开商店上架',
    ]
    for d in diagnosis:
        doc.add_paragraph(d, style='List Bullet')
    
    doc.add_paragraph('解决方案：对54个404包名进行基于包名规律的信息推断（source_site标记为"包名推断"），确保71/71全部入库。')
    
    doc.add_heading('5.5 阶段五：搜索引擎验证', level=2)
    doc.add_paragraph('为消除"包名推断"数据，使用搜索引擎逐一搜索验证：')
    
    add_table(doc,
        ['编号', '包名', '确认结果', '来源'],
        [
            ['1', 'com.ss.android.lark.kaahyz17', 'i讯飞（科大讯飞）', '火鸟手游网'],
            ['2', 'com.ss.android.lark.dagtjt11', '数字国投（国投集团）', '安粉丝网'],
            ['3', 'com.ss.android.lark.sapdl18', '胖东来（胖东来集团）', 'IT168'],
            ['4', 'com.alibaba.taurus.jiangxi', '赣政通（江西省大数据中心）', '当下软件园'],
            ['5', 'com.alibaba.taurus.cpic', '太好钉（太平洋保险）', '安粉丝网'],
            ['6', 'com.alibaba.taurus.qzt', '秦政通（陕西省）', 'IT168'],
            ['7', 'com.alibaba.android.rimet.aliding', '阿里钉（阿里巴巴内部）', 'IT168'],
            ['8', 'com.alibaba.android.rimet.ccflink', 'CCFLink（中国计算机学会）', '极客公园'],
            ['9', 'com.ss.android.lark.htone', '华通3.0（华住集团）', '火鸟手游网'],
            ['10', 'com.alibaba.android.rimet.bgyfw', '碧桂园服务钉钉', '网易新闻（间接确认）'],
            ['11', 'com.alibaba.android.rimet.catlcome', '宁德时代钉钉', '多特软件资讯（间接确认）'],
        ],
        col_widths=[1, 5, 4, 3.5]
    )
    
    doc.add_heading('5.6 阶段六：用户确认与数据纠错', level=2)
    doc.add_paragraph(
        '用户提供完整的71个APP名称列表，结合搜索辅证进行最终确认。此轮重要发现包括：'
    )
    corrections = [
        'com.alibaba.taurus.fujian：原推断"闽政通"→实际"闽政钉"（福建省内部政务协同，非面向市民版）',
        'com.alibaba.taurus.xxxs：原推断"新疆政务钉钉"→实际"学习兴税"（国家税务总局学习平台，xxxs=学习兴税拼音首字母）',
    ]
    for c_item in corrections:
        doc.add_paragraph(c_item, style='List Bullet')
    
    doc.add_paragraph('最终42条"包名推断"记录全部更新为有据可查的验证数据，实现包名推断记录归零。')
    
    doc.add_heading('5.7 阶段七：统计分析报告', level=2)
    doc.add_paragraph(
        '基于验证有效的定制包数据，按飞书、钉钉、专有钉钉、企业微信四个产品线维度，'
        '生成了8章19表的完整统计分析Word报告。报告涵盖行业分布、地域分布、命名规律、'
        '跨平台对比等多维度分析。'
    )

    doc.add_page_break()
    
    # ============================================================
    # 六、数据成果
    # ============================================================
    doc.add_heading('六、数据成果', level=1)
    
    doc.add_heading('6.1 数据库总览', level=2)
    add_table(doc,
        ['数据指标', '数值'],
        [
            ['应用信息库总记录数', '11,247条'],
            ['客户名称库记录数', '673条'],
            ['A级数据（质量评分≥0.8）', '5,334条（47.4%）'],
            ['B级数据（质量评分0.6-0.8）', '91条（0.8%）'],
            ['D级数据（质量评分<0.4）', '5,822条（51.8%）'],
            ['平均质量评分', '0.4086'],
        ],
        col_widths=[5.5, 5.5]
    )
    
    doc.add_paragraph(
        '注：D级数据主要来自滚雪球发现的大量泛办公协同应用，这些记录虽然质量评分较低，'
        '但作为办公协同生态的全景数据仍具有参考价值。核心定制包数据（70个）均已通过多源验证。'
    )
    
    doc.add_heading('6.2 各产品线数据', level=2)
    add_table(doc,
        ['产品线', '总记录数', '占比', '说明'],
        [
            ['办公协同', '10,544', '93.7%', '通过滚雪球发现的广泛办公应用生态'],
            ['企业微信', '566', '5.0%', '企业微信相关应用（含关联应用）'],
            ['钉钉', '63', '0.6%', '钉钉产品线（含企业定制+专有钉钉）'],
            ['飞书', '63', '0.6%', '飞书产品线（含企业定制版）'],
            ['政务服务', '11', '0.1%', '独立政务服务应用'],
        ],
        col_widths=[2.5, 2, 1.5, 8]
    )
    
    doc.add_heading('6.3 定制包统计', level=2)
    doc.add_paragraph('核心成果 —— 共发现并验证70个企业定制安卓客户端包：')
    
    add_table(doc,
        ['产品线', '定制包数量', '占比', '典型案例'],
        [
            ['飞书定制包', '38个', '54.3%', '绿城飞书、海尔飞书、美的飞书、胖东来飞书、数字国投等'],
            ['钉钉企业定制包', '14个', '20.0%', '校园钉、阿里钉、碧桂园服务钉钉、宁德时代钉钉、CCFLink等'],
            ['专有钉钉(政务钉钉)', '15个', '21.4%', '赣政通、闽政钉、秦政通、楚政钉、太好钉(太保)等'],
            ['企业微信', '3个', '4.3%', '标准版、私有版(WeWork@Tencent)、政务版(WeGov)'],
        ],
        col_widths=[3, 2, 1.3, 8]
    )
    
    doc.add_heading('6.4 数据发现方式分布', level=2)
    add_table(doc,
        ['发现方式', '记录数', '说明'],
        [
            ['snowball_round_3', '5,335', '第三轮滚雪球发现'],
            ['snowball_round_2', '4,705', '第二轮滚雪球发现'],
            ['snowball_round_1', '1,025', '第一轮滚雪球发现'],
            ['known_package_verify', '115', '已知种子包名验证'],
            ['user_confirmed_search_assisted', '23', '用户确认+搜索引擎辅证'],
            ['user_confirmed_package_analysis', '17', '用户确认+包名分析'],
            ['search_engine_verified', '11', '搜索引擎直接验证'],
            ['known_*_custom', '14', '已知定制包标识入库'],
            ['news_verified', '2', '新闻报道间接确认'],
        ],
        col_widths=[5, 2, 7]
    )
    
    doc.add_heading('6.5 数据来源分布', level=2)
    add_table(doc,
        ['数据来源', '记录数', '占比'],
        [
            ['应用宝', '11,194', '99.53%'],
            ['用户确认+搜索辅证', '23', '0.20%'],
            ['用户确认+包名分析', '17', '0.15%'],
            ['IT168', '3', '0.03%'],
            ['火鸟手游网', '2', '0.02%'],
            ['新闻报道确认', '2', '0.02%'],
            ['当下软件园', '2', '0.02%'],
            ['安粉丝网', '2', '0.02%'],
            ['极客公园', '1', '0.01%'],
            ['9K9K应用市场', '1', '0.01%'],
        ],
        col_widths=[4, 2, 2]
    )

    doc.add_page_break()
    
    # ============================================================
    # 七、项目文件清单
    # ============================================================
    doc.add_heading('七、项目文件清单', level=1)
    
    doc.add_heading('7.1 代码统计', level=2)
    add_table(doc,
        ['统计项', '数值'],
        [
            ['Python源文件总数', '78个'],
            ['Python代码总行数', '15,982行'],
            ['项目配置文件', '2个（config.py, requirements.txt）'],
            ['需求文档', '1个（pacon_grequire.md）'],
        ],
        col_widths=[5, 5]
    )
    
    doc.add_heading('7.2 核心代码文件', level=2)
    add_table(doc,
        ['文件', '功能说明'],
        [
            ['main.py', '程序入口，三级流水线编排调度器'],
            ['config.py', '全局配置（产品线、搜索引擎、应用商店、反爬参数）'],
            ['crawl_v3.py', '应用宝爬虫v3（核心爬取脚本，含滚雪球发现）'],
            ['crawl_custom_pkgs.py', '定制安卓包专项爬取脚本'],
            ['crawl_404_fallback.py', '404包名多渠道回退爬取脚本'],
            ['update_search_results_v3.py', '搜索引擎验证结果更新脚本（最终版）'],
            ['analyze_custom_pkgs.py', '定制包数据分析脚本'],
            ['generate_report.py', '定制包统计分析Word报告生成脚本'],
            ['known_apps.py', '71个已知定制APP名称映射表'],
        ],
        col_widths=[5.5, 9]
    )
    
    doc.add_heading('7.3 输出文件清单', level=2)
    add_table(doc,
        ['文件名', '大小', '说明'],
        [
            ['results.db', '15.1MB', '应用信息主数据库（SQLite）'],
            ['results.csv', '5.5MB', '应用信息全量CSV导出'],
            ['results.json', '7.6MB', '应用信息全量JSON导出'],
            ['results.xlsx', '1.5MB', '应用信息Excel导出'],
            ['results_飞书.csv', '41.5KB', '飞书产品线数据CSV'],
            ['results_钉钉.csv', '52.1KB', '钉钉产品线数据CSV'],
            ['results_企业微信.csv', '357.5KB', '企业微信产品线数据CSV'],
            ['results_办公协同.csv', '5.1MB', '办公协同产品线数据CSV'],
            ['results_政务服务.csv', '7.5KB', '政务服务产品线数据CSV'],
            ['customers.db', '124.0KB', '企业客户名称库（SQLite）'],
            ['customers.csv', '61.5KB', '客户名称CSV导出'],
            ['customers.json', '174.4KB', '客户名称JSON导出'],
            ['analysis_data.json', '16.3KB', '定制包分析中间数据'],
            ['report.txt', '1.4KB', '简要统计报告'],
            ['report.json', '1.2KB', '统计报告JSON'],
            ['企业办公协同应用定制包统计分析报告.docx', '50.5KB', '定制包详细统计分析报告（8章19表）'],
        ],
        col_widths=[5.5, 1.5, 7]
    )

    doc.add_page_break()
    
    # ============================================================
    # 八、问题与解决方案
    # ============================================================
    doc.add_heading('八、问题与解决方案', level=1)
    
    add_table(doc,
        ['问题编号', '问题描述', '解决方案', '效果'],
        [
            ['P-001', '应用宝限流导致v2仅获得27条记录', '降频(0.5-1.2s)、重试(MAX_RETRY=2)、滚雪球发现', '提升至11,247条'],
            ['P-002', '54个企业定制包在所有公开商店404', '包名规律推断→搜索引擎验证→用户确认三轮迭代', '71/71全部入库并验证'],
            ['P-003', '小米应用商店不返回404而给通用首页', 'fix_xiaomi_bad_data.py清理无效数据', '54条无效数据已修复'],
            ['P-004', '包名推断数据可靠性不足', '搜索引擎直接验证+新闻间接确认+用户提供名称', '推断记录归零'],
            ['P-005', '政务钉钉面向市民版与内部版包名混淆', '搜索引擎验证区分（如闽政通vs闽政钉）', '2条关键纠错'],
            ['P-006', '编码式包名（da*/ka*/sa*）无法公开查找', '结合用户提供的企业名称确认', '17条编码式包名全部确认'],
        ],
        col_widths=[1.5, 4, 5, 3.5]
    )

    doc.add_page_break()
    
    # ============================================================
    # 九、项目总结
    # ============================================================
    doc.add_heading('九、项目总结', level=1)
    
    doc.add_heading('9.1 项目成果总结', level=2)
    doc.add_paragraph('本项目成功完成了企业IM安卓客户端包信息的自动化采集系统开发，主要成果如下：')
    
    achievements = [
        '构建了完整的三级递进式爬虫系统：78个Python模块，15,982行代码，覆盖8+应用商店',
        '建立了包含11,247条记录的应用信息数据库，覆盖企业微信、钉钉、飞书及办公协同生态',
        '发现并验证了70个企业定制安卓客户端包，其中飞书38个、钉钉14个、专有钉钉15个、企业微信3个',
        '通过应用宝爬取+搜索引擎验证+新闻确认+用户信息四重手段，实现了核心数据零推断、全验证',
        '产出了完整的多格式数据导出（SQLite/CSV/JSON/XLSX）和专业的统计分析Word报告',
        '建立了673条企业客户名称库，为后续持续采集提供了基础数据',
    ]
    for a in achievements:
        doc.add_paragraph(a, style='List Bullet')
    
    doc.add_heading('9.2 关键发现', level=2)
    findings = [
        '飞书在定制包数量上领先（38个），显示其"大客户深度定制"策略，覆盖制造/工业/地产等行业',
        '钉钉+专有钉钉合计29个定制包，形成"企业+政务"双轮驱动模式，政务覆盖9+省级行政区',
        '企业微信仅3个包名版本，体现"一个平台、多种部署"的产品策略',
        '大量企业定制包仅通过企业内部MDM分发，不在任何公开应用商店上架，这是信息采集的主要挑战',
        '编码式包名（如da*/ka*/sa*前缀）是飞书的特有现象，反映其标准化的企业定制流程',
        '政务钉钉的面向市民端APP与内部协同APP使用不同包名，需要区分对待',
    ]
    for f in findings:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_heading('9.3 不足与改进方向', level=2)
    improvements = [
        '当前仅以应用宝为主要数据源，华为/小米/OPPO/vivo等商店的爬虫可进一步完善',
        '滚雪球发现的D级数据（5,822条）质量较低，可通过更精细的过滤策略提升整体数据质量',
        '搜索引擎验证环节依赖人工辅助较多，后续可考虑引入自动化NLP提取和验证',
        '数据更新机制尚未实现，需增加增量爬取和定期刷新功能',
        '反爬策略可进一步增强，如接入付费代理池、增加IP轮换',
    ]
    for imp in improvements:
        doc.add_paragraph(imp, style='List Bullet')
    
    doc.add_heading('9.4 项目数据汇总', level=2)
    
    # 最终汇总表
    add_table(doc,
        ['维度', '数据'],
        [
            ['项目开发周期', '2026年3月9日（集中开发）'],
            ['代码规模', '78个Python文件 / 15,982行代码'],
            ['模块数量', '7个功能模块（3级爬虫+管道+存储+工具+业务脚本）'],
            ['应用信息库', '11,247条记录 / 15.1MB'],
            ['客户名称库', '673条记录 / 124KB'],
            ['定制包总数', '70个（飞书38+钉钉14+专有钉钉15+企微3）'],
            ['数据验证率', '核心定制包100%验证（搜索引擎+新闻+用户确认）'],
            ['输出文件', '18个文件（DB/CSV/JSON/XLSX/DOCX），总计约42MB'],
            ['分析报告', '8章19表的定制包统计分析报告'],
            ['需求文档', 'pacon_grequire.md（487行，完整的开发需求规格）'],
        ],
        col_widths=[3, 11]
    )
    
    # ============================================================
    # 文档信息
    # ============================================================
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line_p.add_run('━' * 40)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f'报告生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer2.add_run('— 报告结束 —')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    # 保存文档
    output_path = os.path.join('output', '企业IM安卓客户端包信息爬虫-项目开发报告.docx')
    doc.save(output_path)
    print(f'✅ 项目开发报告已生成: {output_path}')
    print(f'   文件大小: {os.path.getsize(output_path)/1024:.1f} KB')
    
    return output_path


if __name__ == '__main__':
    generate_report()
